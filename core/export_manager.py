"""小说导出与完整项目备份。"""
from __future__ import annotations

import html
import json
import zipfile
from datetime import datetime
from pathlib import Path

from filelock import FileLock


class ExportManager:
    def __init__(self, novel_name: str, novel_path: Path, logger):
        self.name = novel_name
        self.path = novel_path
        self.logger = logger
        self.output = novel_path / "exports"
        self.output.mkdir(parents=True, exist_ok=True)
        titles_file = novel_path / "outline" / "chapter_titles.json"
        try:
            self.titles = json.loads(titles_file.read_text("utf-8")) if titles_file.exists() else {}
        except Exception:
            self.titles = {}

    def _chapter_title(self, chapter: int) -> str:
        title = str(self.titles.get(str(chapter), "")).strip()
        return f"第{chapter}章 {title}" if title else f"第{chapter}章"

    def _chapters(self) -> list[tuple[int, str]]:
        result = []
        for file in sorted((self.path / "chapters").glob("*.txt")):
            try:
                result.append((int(file.stem), file.read_text("utf-8", errors="replace")))
            except (ValueError, OSError):
                continue
        return result

    def export(self, format_name: str) -> Path:
        format_name = format_name.lower()
        handlers = {"txt": self._txt, "md": self._markdown, "docx": self._docx, "epub": self._epub, "zip": self._zip}
        if format_name not in handlers:
            raise ValueError(f"不支持的导出格式: {format_name}")
        path = handlers[format_name]()
        self.logger.info("导出小说: %s -> %s", self.name, path)
        return path

    def _txt(self) -> Path:
        target = self.output / f"{self.name}.txt"
        parts = [self.name, "=" * len(self.name), ""]
        for chapter, content in self._chapters():
            parts.extend([self._chapter_title(chapter), "", content.strip(), ""])
        target.write_text("\n".join(parts), "utf-8")
        return target

    def _markdown(self) -> Path:
        target = self.output / f"{self.name}.md"
        parts = [f"# {self.name}", ""]
        for chapter, content in self._chapters():
            parts.extend([f"## {self._chapter_title(chapter)}", "", content.strip(), ""])
        target.write_text("\n".join(parts), "utf-8")
        return target

    def _docx(self) -> Path:
        from docx import Document
        document = Document()
        document.add_heading(self.name, 0)
        for chapter, content in self._chapters():
            document.add_heading(self._chapter_title(chapter), level=1)
            for paragraph in content.splitlines():
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())
        target = self.output / f"{self.name}.docx"
        document.save(target)
        return target

    def _epub(self) -> Path:
        from ebooklib import epub
        book = epub.EpubBook()
        book.set_identifier(f"novel-{self.name}-{datetime.now().strftime('%Y%m%d')}")
        book.set_title(self.name)
        book.set_language("zh-CN")
        book.add_author("本地AI小说工作台")
        items = []
        for chapter, content in self._chapters():
            title = self._chapter_title(chapter)
            item = epub.EpubHtml(title=title, file_name=f"chapter_{chapter:06d}.xhtml", lang="zh-CN")
            paragraphs = "".join(f"<p>{html.escape(line)}</p>" for line in content.splitlines() if line.strip())
            item.content = f"<h1>{html.escape(title)}</h1>{paragraphs}"
            book.add_item(item)
            items.append(item)
        book.toc = tuple(items)
        book.spine = ["nav", *items]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        target = self.output / f"{self.name}.epub"
        epub.write_epub(str(target), book)
        return target

    def _zip(self) -> Path:
        target = self.output / f"{self.name}_完整项目_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        temporary = target.with_suffix(".zip.tmp")
        try:
            with FileLock(str(self.path / ".novel_mutation.lock"), timeout=600):
                with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as archive:
                    for file in self.path.rglob("*"):
                        if (
                            not file.is_file()
                            or self.output in file.parents
                            or file.name.endswith((".lock", ".tmp"))
                        ):
                            continue
                        archive.write(file, file.relative_to(self.path))
                temporary.replace(target)
        except Exception:
            temporary.unlink(missing_ok=True)
            raise
        return target
